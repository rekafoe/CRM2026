# -*- coding: utf-8 -*-
"""Генерация ответа на заявку в/ч 1463 в формате .docx"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

OUT = Path(__file__).resolve().parent / "KP_otvet_vch1463.docx"


def main() -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(1.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Кому: Руководителю (по прилагаемому перечню)")
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Заказчик: Войсковая часть 1463, 222712, Минская обл., г. Дзержинск, ул. Чкалова, 30, УНП 600005411").font.size = Pt(10)

    doc.add_paragraph()
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("Исх. № _____ от «___» __________ 2026 г.")
    r.bold = True
    r.font.size = Pt(12)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "О предоставлении коммерческого предложения по заявке от 10.03.2026 г. № ___192___ "
        "о закупке из одного источника (ознакомление с информацией, изучение конъюнктуры рынка)"
    )
    r.font.size = Pt(11)

    doc.add_paragraph()
    body = doc.add_paragraph()
    body.add_run("Уважаемый коллега!").font.size = Pt(12)

    intro = doc.add_paragraph()
    intro.paragraph_format.first_line_indent = Cm(1.25)
    intro.add_run(
        "Направляем коммерческое предложение в соответствии с требованиями заявки и таблицей "
        "предмета закупки (информационные таблички, ОКРБ 22.29.29.500)."
    ).font.size = Pt(12)

    doc.add_paragraph()
    t1 = doc.add_paragraph()
    t1.add_run("1. Сведения об организации").bold = True
    t1.runs[0].font.size = Pt(12)

    table1 = doc.add_table(rows=5, cols=2)
    table1.style = "Table Grid"
    rows1 = [
        ("Полное наименование", "[ООО «…»]"),
        ("Юридический адрес", "[индекс, город, улица, дом, офис]"),
        ("УНП", "[____________]"),
        ("Контактное лицо", "[ФИО полностью]"),
        ("Телефон, e-mail", "[+375 __ ___-__-__], [e-mail]"),
    ]
    for i, (a, b) in enumerate(rows1):
        table1.rows[i].cells[0].text = a
        table1.rows[i].cells[1].text = b
    for row in table1.rows:
        for cell in row.cells:
            for par in cell.paragraphs:
                for run in par.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()
    t2h = doc.add_paragraph()
    t2h.add_run("2. Коммерческое предложение (цены с НДС)").bold = True
    t2h.runs[0].font.size = Pt(12)

    cond = doc.add_paragraph()
    cond.paragraph_format.first_line_indent = Cm(1.25)
    cond.add_run(
        "Предлагаем поставку товара нового (не бывшего в употреблении, без восстановления и замены "
        "составных частей с восстановлением потребительских свойств в значении, указанном в заявке), "
        "с гарантией не менее 12 месяцев с даты передачи заказчику при соблюдении условий эксплуатации. "
        "Материал основы: ПВХ, толщина не менее 4 мм; изображение: УФ-печать; макеты: согласно Приложению № 1 к заявке "
        "и согласованным файлам сторон."
    ).font.size = Pt(11)

    table2 = doc.add_table(rows=6, cols=5)
    table2.style = "Table Grid"
    hdr = ["№ п/п", "Наименование (требования по заявке)", "Объём", "Цена за 1 ед. с НДС, BYN", "Сумма с НДС, BYN"]
    for j, h in enumerate(hdr):
        table2.rows[0].cells[j].text = h
    data = [
        (
            "1",
            "Информационная табличка 10×25 см, ПВХ не менее 4 мм, УФ-печать (по Приложению № 1)",
            "12 шт.",
            "5,01",
            "60,14",
        ),
        (
            "2",
            "Информационная табличка 20×30 см, ПВХ не менее 4 мм, УФ-печать (по Приложению № 1)",
            "5 шт.",
            "8,07",
            "40,37",
        ),
        (
            "3",
            "Информационная табличка 25×40 см с карманом ПВХ, УФ-печать (по Приложению № 1)",
            "1 шт.",
            "16,68",
            "16,68",
        ),
        (
            "4",
            "Информационная табличка 105×160 см с бортом 3 см, карман ПВХ 90×150 см, УФ-печать; "
            "при необходимости карта на плотной бумаге в комплекте — по согласованному макету Приложения № 1",
            "5 шт.",
            "182,70",
            "913,49",
        ),
    ]
    for i, row_data in enumerate(data, start=1):
        for j, val in enumerate(row_data):
            table2.rows[i].cells[j].text = val

    for row in table2.rows:
        for cell in row.cells:
            for par in cell.paragraphs:
                for run in par.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()
    tot = doc.add_paragraph()
    tot.add_run(
        "Всего наименований: 23 шт. Итого стоимость с НДС: 1 030,68 BYN "
        "(одна тысяча тридцать белорусских рублей 68 копеек). "
        "Стоимость не превышает предельную цену закупки 2 000,00 руб. по заявке."
    ).font.size = Pt(11)

    pay = doc.add_paragraph()
    pay.paragraph_format.first_line_indent = Cm(1.25)
    pay.add_run(
        "Условия оплаты: в соответствии с заявкой — через органы государственного казначейства "
        "(реквизиты и счёт выставляются после согласования заказчиком)."
    ).font.size = Pt(11)

    deliv = doc.add_paragraph()
    deliv.paragraph_format.first_line_indent = Cm(1.25)
    deliv.add_run(
        "Срок и место поставки: с 18.05.2026 по 29.05.2026, грузополучатель: в/ч 1463, г. Дзержинск, ул. Чкалова, 30 "
        "(порядок приёмки уточняется по контактам заказчика)."
    ).font.size = Pt(11)

    ready = doc.add_paragraph()
    ready.paragraph_format.first_line_indent = Cm(1.25)
    ready.add_run(
        "Готовы по запросу направить счёт, реквизиты, уточнить финальные файлы для производства и спецификацию."
    ).font.size = Pt(11)

    doc.add_paragraph()
    note = doc.add_paragraph()
    r = note.add_run("Замечания перед отправкой:\n")
    r.bold = True
    r.font.size = Pt(10)
    note.add_run(
        "• Конечный срок ответа на заявку в документе заказчика: 5 мая 2026 г. до 14:00 — при необходимости "
        "согласуйте приём ответа с заказчиком.\n"
        "• Каналы: e-mail hudaev_as@ops.gov.by, факс 8 (017) 555-31-58, 555-31-56, почта: 222712, "
        "Минская обл., г. Дзержинск, ул. Чкалова, 30.\n"
        "• Укажите реальные полное наименование, адрес и УНП вашей организации."
    ).font.size = Pt(10)

    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.add_run("Руководитель _________________ / _________________ /    М.П.").font.size = Pt(11)

    doc.save(OUT)
    print("Saved:", OUT)


if __name__ == "__main__":
    main()
